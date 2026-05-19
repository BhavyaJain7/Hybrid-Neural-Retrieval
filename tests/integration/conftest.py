"""
Integration test shared fixtures.

Philosophy
----------
Integration tests sit one layer above unit tests:
  - Real file system (tmp_path from pytest)
  - Real Settings object pointing at tmp dirs
  - Real CollectionManager, BM25sRetriever, QdrantRetriever in-process
  - Only true external boundaries are mocked:
      • Groq (network + paid API)
      • SentenceTransformer (heavy download, slow)

This means we catch real bugs in the wiring between components
(parser → chunker → BM25 → Qdrant → API) without hitting the internet.
"""
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch


# ── Real Settings pointed at tmp dirs ─────────────────────────────────────────
@pytest.fixture(autouse=True)
def real_settings(tmp_path, monkeypatch):
    """
    Creates a real Settings instance wired to isolated tmp dirs.
    Overwrites the global `settings` singleton in every module so there
    are no leftover cross-test side-effects.
    """
    from neural_search import config
    from neural_search.config import Settings

    # Build a real Settings with tmp paths — no .env needed
    s = Settings(
        groq_api_key="test-groq-key",
        groq_model="llama-3.1-8b-instant",
        embedding_model="all-MiniLM-L6-v2",
        chunk_size=128,
        chunk_overlap=16,
        top_k=5,
        rrf_k=60,
        data_dir=tmp_path / "data",
        qdrant_path=tmp_path / "data" / "qdrant",
        bm25_index_path=tmp_path / "data" / "bm25",
    )
    s.ensure_dirs()
    monkeypatch.setattr(config, "settings", s)

    # Patch module-level `settings` in modules that expose it directly.
    # NOTE: pipeline.py uses `import settings as global_settings` (not `settings`),
    # so it is intentionally excluded — use settings_obj= kwarg for injection there.
    _modules_with_settings = [
        "neural_search.retrieval.dense",
        "neural_search.retrieval.sparse",
        "neural_search.retrieval.hybrid",
        "neural_search.collections.manager",
        "neural_search.api.routes",
    ]
    import sys
    for mod_name in _modules_with_settings:
        if mod_name in sys.modules:
            mod = sys.modules[mod_name]
            if hasattr(mod, "settings"):
                monkeypatch.setattr(mod, "settings", s)

    return s


# ── Fake SentenceTransformer (deterministic, fast) ───────────────────────────
@pytest.fixture
def fake_embedding_model():
    """
    Returns a mock SentenceTransformer whose outputs are deterministic
    numpy vectors. This prevents downloading the real model in CI and
    makes tests fast while still exercising all surrounding code.
    """
    import numpy as np

    model = MagicMock()
    model.get_sentence_embedding_dimension.return_value = 384

    def _encode(texts, **kwargs):
        if isinstance(texts, str):
            rng = sum(ord(c) for c in texts) % 1000
            vec = np.array([rng / 1000.0] * 384, dtype=np.float32)
            return vec / (np.linalg.norm(vec) or 1.0)
        out = []
        for t in texts:
            rng = sum(ord(c) for c in t) % 1000
            vec = np.array([rng / 1000.0] * 384, dtype=np.float32)
            vec = vec / (np.linalg.norm(vec) or 1.0)
            out.append(vec)
        return np.array(out, dtype=np.float32)

    model.encode.side_effect = _encode
    return model


# ── Minimal real PDF ──────────────────────────────────────────────────────────
@pytest.fixture
def small_pdf(tmp_path) -> Path:
    """Creates a 3-page PDF with enough text to produce at least 1 chunk."""
    try:
        import fitz
    except ImportError:
        pytest.skip("pymupdf not installed")

    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        # Insert enough words to form a real chunk
        content = (
            f"Page {i + 1}: This is a test document about neural search systems. "
            "Hybrid retrieval combines sparse BM25 ranking with dense vector similarity. "
            "Reciprocal Rank Fusion merges results from both retrievers deterministically. "
            "Documents are parsed, chunked, and indexed into both retrieval backends. "
            "This sentence adds more tokens to ensure the chunk threshold is met. "
        )
        page.insert_text((50, 72), content)
    path = tmp_path / "test_doc.pdf"
    doc.save(str(path))
    doc.close()
    return path


# ── Minimal real DOCX ─────────────────────────────────────────────────────────
@pytest.fixture
def small_docx(tmp_path) -> Path:
    """Creates a 2-section DOCX with enough text for at least 1 chunk."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This is a test DOCX document for integration testing of the neural search pipeline. "
        "BM25 and dense retrieval are combined via Reciprocal Rank Fusion. "
        "Each section provides enough tokens to form a valid chunk for indexing. "
    )
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "The ingestion pipeline parses documents into pages, chunks them, "
        "and indexes them into both the BM25 sparse index and the Qdrant vector store. "
        "This method section adds additional content to test multi-chunk documents. "
    )
    path = tmp_path / "test_doc.docx"
    doc.save(str(path))
    return path
