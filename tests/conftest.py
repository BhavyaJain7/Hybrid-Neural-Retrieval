"""
Shared fixtures for all tests.
Design rules:
- patch_settings is autouse — every test gets a safe Settings mock, no .env needed
- All mock attributes match the actual Settings schema exactly
- No local mock_settings fixtures in individual test files — use this one
- sample_chunks, sample_pdf, sample_docx available to all tests
"""
import pytest
from pathlib import Path
from unittest.mock import MagicMock, create_autospec


# ── Autouse: patch settings before every test ─────────────────────────────────

@pytest.fixture(autouse=True)
def patch_settings(tmp_path, monkeypatch):
    """
    Replaces live settings with a spec-constrained mock matching Settings exactly.
    autouse=True means this runs for every test automatically.
    """
    from neural_search import config
    from neural_search.config import Settings

    mock = create_autospec(Settings, instance=True)

    # Mirror exact field names from Settings schema
    mock.groq_api_key = "test-key"
    mock.groq_model = "llama3-8b-8192"
    mock.embedding_model = "all-MiniLM-L6-v2"
    mock.reranker_model = "cross-encoder/ms-marco-MiniLM-L-6-v2"  # Phase 4
    mock.chunk_size = 256
    mock.chunk_overlap = 32
    mock.top_k = 5
    mock.rrf_k = 60
    mock.data_dir = tmp_path / "data"
    mock.qdrant_path = tmp_path / "data" / "qdrant"
    mock.bm25_index_path = tmp_path / "data" / "bm25_index"
    mock.bm25_path_for.side_effect = lambda slug: tmp_path / "data" / "bm25_index" / slug
    mock.documents_path_for.side_effect = lambda slug: tmp_path / "data" / "documents" / slug
    mock.snapshot_path_for.side_effect = lambda slug: tmp_path / "data" / "snapshots" / slug
    mock.ensure_dirs.return_value = None
    mock.assert_groq_configured.return_value = None

    monkeypatch.setattr(config, "settings", mock)
    return mock


# ── Sample chunks ─────────────────────────────────────────────────────────────

@pytest.fixture
def sample_chunks():
    from neural_search.ingestion.chunker import Chunk
    return [
        Chunk(
            chunk_id=f"chunk_{i:04d}",
            doc_id="test_doc",
            source_file="test.pdf",
            page=i + 1,
            chunk_index=i,
            text=f"Sample text for chunk {i}. Contains content about topic {i}.",
            token_count=15 + i,
        )
        for i in range(10)
    ]


# ── Real PDF fixture ──────────────────────────────────────────────────────────

@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    try:
        import fitz
    except ImportError:
        pytest.skip("pymupdf not installed")

    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        page.insert_text(
            (50, 50),
            f"Page {i + 1} content. Test document text for neural search.",
        )
    path = tmp_path / "sample.pdf"
    doc.save(str(path))
    doc.close()
    return path


# ── Real DOCX fixture ─────────────────────────────────────────────────────────

@pytest.fixture
def sample_docx(tmp_path) -> Path:
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Introduction section with test content for neural search.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Methods section describing the approach used.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path
